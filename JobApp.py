import os
import logging
import time
from flask import Flask, render_template, request, jsonify, send_file, session
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pdfplumber
from groq import Groq
from io import BytesIO

app = Flask(__name__)
app.secret_key = "your_secret_key"

# Configure logging
logging.basicConfig(
    filename="job_application.log",
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s"
)

# Initialize Groq client
groq_client = Groq(api_key="paste groq cloud API")

# Helper function to handle API calls with exponential backoff
def call_groq_with_backoff(api_func, *args, **kwargs):
    max_retries = 5
    delay = 1  # Initial delay in seconds
    for attempt in range(max_retries):
        try:
            response = api_func(*args, **kwargs)
            return response
        except Exception as e:
            if "429 Too Many Requests" in str(e):
                logging.info(f"Rate limit hit, retrying in {delay} seconds...")
                time.sleep(delay)
                delay *= 2  # Exponential backoff
            else:
                logging.error(f"API call failed: {e}")
                raise
    raise RuntimeError("Max retries exceeded for API call.")

# Helper function to read resume content
def read_resume(file_path):
    if file_path.endswith('.pdf'):
        with pdfplumber.open(file_path) as pdf:
            return "\n".join(page.extract_text() for page in pdf.pages)
    elif file_path.endswith('.docx'):
        doc = Document(file_path)
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)
    else:
        raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")

# Helper function to optimize resume for ATS
def optimize_resume_for_ats(resume_text, job_description, extra_info=None, target_score=80):
    optimized_resume = resume_text
    ats_score = 0

    for attempt in range(3):  # Limit attempts to avoid excessive API calls
        prompt = (
            f"Rewrite the following resume to align with this job description for ATS-friendly formatting and maximize the ATS score.\n\n"
            f"Resume:\n{optimized_resume}\n\n"
            f"Job Description:\n{job_description}\n\n"
            f"Extra Instructions: {extra_info if extra_info else ''}\n"
        )
        response = call_groq_with_backoff(
            groq_client.chat.completions.create,
            messages=[{"role": "user", "content": prompt}],
            model="llama3-8b-8192"
        )
        optimized_resume = response.choices[0].message.content.strip()
        ats_score = calculate_ats_score(optimized_resume, job_description)
        logging.info(f"Attempt {attempt + 1}: Updated ATS Score = {ats_score}")

        if ats_score >= target_score:
            break

    return optimized_resume

# Helper function to generate cover letter
def generate_cover_letter(resume_text, job_description, extra_info=None):
    prompt = (
        f"Write a professional cover letter based on the following resume and job description.\n\n"
        f"Resume:\n{resume_text}\n\n"
        f"Job Description:\n{job_description}\n\n"
        f"Extra Instructions: {extra_info if extra_info else ''}\n"
    )
    response = call_groq_with_backoff(
        groq_client.chat.completions.create,
        messages=[{"role": "user", "content": prompt}],
        model="llama3-8b-8192"
    )
    return response.choices[0].message.content.strip()

# Helper function to calculate ATS score
def calculate_ats_score(resume_text, job_description):
    score = sum(1 for word in job_description.split() if word.lower() in resume_text.lower())
    total_words = len(job_description.split())
    return round((score / total_words) * 100, 2)

# Save content into Harvard CV Template
def save_to_template(content, template_path, output_path):
    doc = Document(template_path)

    # Clear placeholders if present
    for paragraph in doc.paragraphs:
        paragraph.text = ""  # Clear existing content

    # Populate the template
    for line in content.split("\n"):
        paragraph = doc.add_paragraph(line)
        paragraph.style.font.name = 'Times New Roman'
        paragraph.style.font.size = Pt(12)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.save(output_path)
    return output_path

@app.route("/", methods=["GET", "POST"])
def home():
    if request.method == "POST":
        resume_file = request.files.get("resume")
        job_description = request.form.get("job_description")
        extra_info = request.form.get("extra_info", "")

        if not resume_file or not job_description:
            return jsonify({"error": "Please upload a resume and provide a job description."})

        os.makedirs("uploads", exist_ok=True)
        resume_path = os.path.join("uploads", resume_file.filename)
        resume_file.save(resume_path)

        try:
            resume_text = read_resume(resume_path)
            optimized_resume = optimize_resume_for_ats(resume_text, job_description, extra_info)
            optimized_cover_letter = generate_cover_letter(resume_text, job_description, extra_info)
            ats_score = calculate_ats_score(optimized_resume, job_description)

            resume_output_path = "/Users/vivekkumar/AI/uploads/optimized_resume.docx"
            cover_letter_output_path = "/Users/vivekkumar/AI/uploads/optimized_cover_letter.docx"

            save_to_template(optimized_resume, "/Users/vivekkumar/AI/templates/MAKE A COPY OF THE DOC - Harvard CV Template.docx", resume_output_path)
            save_to_template(optimized_cover_letter, "/Users/vivekkumar/AI/templates/MAKE A COPY OF THE DOC - Harvard CV Template.docx", cover_letter_output_path)

            session['resume_output_path'] = resume_output_path
            session['cover_letter_output_path'] = cover_letter_output_path

            return jsonify({
                "resume": optimized_resume,
                "cover_letter": optimized_cover_letter,
                "ats_score": ats_score,
                "download_resume_url": "/download_resume",
                "download_cover_letter_url": "/download_cover_letter"
            })
        except Exception as e:
            logging.error(f"Error processing application: {e}")
            return jsonify({"error": "An error occurred while generating the documents. Please check the logs."})

    return render_template("index.html")

@app.route("/download_resume")
def download_resume():
    resume_path = session.get('resume_output_path')
    if resume_path and os.path.exists(resume_path):
        return send_file(resume_path, as_attachment=True, download_name="optimized_resume.docx")
    return jsonify({"error": "No resume available for download."})

@app.route("/download_cover_letter")
def download_cover_letter():
    cover_letter_path = session.get('cover_letter_output_path')
    if cover_letter_path and os.path.exists(cover_letter_path):
        return send_file(cover_letter_path, as_attachment=True, download_name="optimized_cover_letter.docx")
    return jsonify({"error": "No cover letter available for download."})

if __name__ == "__main__":
    app.run(debug=True)
